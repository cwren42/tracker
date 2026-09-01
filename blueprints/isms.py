import html
import io
import re
from difflib import HtmlDiff, SequenceMatcher
from functools import partial
from pathlib import Path

from datetime import datetime

from flask import Blueprint, abort, flash, redirect, render_template, request, send_file, url_for
from flask_login import login_required, current_user
import markdown as markdown_lib
from openpyxl import load_workbook
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.pdfgen import canvas
from reportlab.platypus import (
    HRFlowable, Image, PageBreak, Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle,
)

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, Inches, RGBColor
from werkzeug.utils import secure_filename

from extensions import db
from models import Asset, AuditTrail, ISMSDocument, ISMSDocumentVersion, ISMSExportRun
from utils import admin_required
import isms_ledger_service


bp = Blueprint('isms', __name__)


GLOCALIZATION_SURVEY_TEMPLATE_PATH = Path(
    '/var/www/tracker/templates/ISMS-MANUAL/_Temp-Files/Glocalization-Survey-Cirque-Filled-V7-2026-05-08.xlsx'
)

GLOCALIZATION_SURVEY_OVERRIDES = {
    43: {
        'status': 'Exist & No corrections',
        'reason': 'CSIRT function covered by Cirque Incident Management Policy and Incident Response Procedures, including the severity matrix and breach notification timelines.',
    },
    45: {
        'status': 'Exist & No corrections',
        'reason': 'Backup and recovery covered by Cirque Backup and Restoration Procedure with retention and recovery objectives defined in the current manual set.',
    },
    71: {
        'status': 'Exist & No corrections',
        'reason': 'Incident handling covered by Cirque Incident Response Procedures (Global, US, and Asia localized), including severity classification and jurisdiction-specific breach clocks.',
    },
}


def _current_actor():
    return getattr(current_user, 'display_name', None) or getattr(current_user, 'username', None) or 'system'


def _current_actor_username():
    return getattr(current_user, 'username', None) or _current_actor()


def _get_current_version(document):
    version = document.current_version
    if version is None and document.versions:
        version = document.versions[0]
    return version


def _create_version(document, markdown_body, change_summary, *, is_restore=False, restored_from=None):
    latest_version = _get_current_version(document)
    next_version_number = 1 if latest_version is None else latest_version.version_number + 1
    new_version = ISMSDocumentVersion(
        document_id=document.id,
        version_number=next_version_number,
        markdown_body=markdown_body,
        rendered_html=render_markdown(markdown_body),
        change_summary=change_summary,
        is_restore=is_restore,
        restored_from_version_id=restored_from.id if restored_from else None,
        created_by=_current_actor_username(),
    )
    db.session.add(new_version)
    db.session.flush()

    document.current_version = new_version
    document.updated_by = _current_actor_username()
    document.updated_at = datetime.utcnow()
    return new_version


def _log_action(document, action, details):
    try:
        db.session.add(
            AuditTrail(
                asset_id=None,
                action=action,
                table_name='isms_document',
                record_id=document.id,
                old_values=None,
                new_values=details,
                changed_by=_current_actor_username(),
            )
        )
    except Exception:
        pass


def _download_filename(document, version_number, extension):
    return f"{document.slug or 'isms-document'}_v{version_number}.{extension}"


def _slugify_heading(text_value, seen_slugs=None):
    slug = re.sub(r'[^a-z0-9]+', '-', text_value.lower()).strip('-') or 'section'
    if seen_slugs is None:
        return slug

    base_slug = slug
    suffix = 2
    while slug in seen_slugs:
        slug = f'{base_slug}-{suffix}'
        suffix += 1
    seen_slugs.add(slug)
    return slug


def _flatten_toc_tokens(tokens, level=1):
    items = []
    for token in tokens or []:
        items.append({
            'id': token.get('id'),
            'title': token.get('name', ''),
            'level': level,
        })
        if token.get('children'):
            items.extend(_flatten_toc_tokens(token['children'], level + 1))
    return items


def _normalize_markdown_source(markdown_body):
    normalized = re.sub(r'^---\s*\n.*?\n---\s*\n?', '', markdown_body, flags=re.DOTALL)
    normalized = re.sub(r'^\\newpage\s*$', '', normalized, flags=re.MULTILINE)
    normalized = re.sub(r'\\([`*_{}\[\]()#+\-.!|])', r'\1', normalized)
    return normalized.strip()


def _extract_metadata_fields(line):
    fields = []
    cursor = 0

    while cursor < len(line):
        start = line.find('**', cursor)
        if start == -1:
            break
        end = line.find('**', start + 2)
        if end == -1:
            return []

        token = line[start + 2:end].strip()
        cursor = end + 2
        if ':' in token:
            label, inline_value = token.split(':', 1)
        else:
            label = token
            inline_value = ''
            if cursor >= len(line) or line[cursor] != ':':
                return []
            cursor += 1

        next_start = line.find('**', cursor)
        trailing_value = line[cursor:next_start if next_start != -1 else len(line)].strip()
        value = ' '.join(part for part in [inline_value.strip(), trailing_value] if part)
        fields.append((label.strip(), value.strip()))

        if next_start == -1:
            break
        cursor = next_start

    return fields


def _normalize_metadata_fields(fields, heading_text):
    normalized_fields = []
    normalized_heading = re.sub(r'^IS-[A-Z0-9-]+:\s*', '', heading_text).strip().lower()

    label_map = {
        'Document': 'Document ID',
        'Standards Name': 'Document Title',
        'Standard Type': 'Document Type',
    }
    suppressed_labels = {'Standard Retention', 'Document ID'}

    for label, value in fields:
        label = label_map.get(label, label)
        if label in suppressed_labels:
            continue

        normalized_value = value.strip()
        if label == 'Document Title' and normalized_value.lower() == normalized_heading:
            continue
        normalized_fields.append((label, normalized_value))

    return normalized_fields


def _build_metadata_block(fields, heading_text):
    normalized_fields = _normalize_metadata_fields(fields, heading_text)
    items = []
    for label, value in normalized_fields:
        items.append(
            '<div class="isms-section-meta-item">'
            f'<div class="isms-section-meta-label">{html.escape(label)}</div>'
            f'<div class="isms-section-meta-value">{_render_inline(value)}</div>'
            '</div>'
        )
    return '<div class="isms-section-meta">' + ''.join(items) + '</div>'


def _prepare_reader_markdown(markdown_body):
    normalized = _normalize_markdown_source(markdown_body)
    lines = normalized.split('\n')
    prepared = []
    index = 0

    while index < len(lines):
        line = lines[index]
        heading_match = re.match(r'^(#{1,6})\s+(.+?)\s*$', line)
        prepared.append(line)
        index += 1

        if not heading_match:
            continue

        heading_text = heading_match.group(2).strip()

        while index < len(lines) and not lines[index].strip():
            prepared.append(lines[index])
            index += 1

        while index < len(lines):
            duplicate_heading = re.match(r'^(#{1,6})\s+(.+?)\s*$', lines[index].strip())
            if not duplicate_heading or duplicate_heading.group(2).strip() != heading_text:
                break
            index += 1
            while index < len(lines) and not lines[index].strip():
                index += 1

        if index < len(lines):
            duplicate_match = re.match(r'^\*\*(.+?)\*\*\s*$', lines[index].strip())
            if duplicate_match and duplicate_match.group(1).strip() == heading_text:
                index += 1
                while index < len(lines) and not lines[index].strip():
                    index += 1

        metadata_fields = []
        metadata_index = index
        while metadata_index < len(lines):
            candidate = lines[metadata_index].strip()
            if not candidate:
                break
            if re.match(r'^#{1,6}\s+', candidate):
                break
            parsed_fields = _extract_metadata_fields(candidate)
            if not parsed_fields:
                break
            metadata_fields.extend(parsed_fields)
            metadata_index += 1

        if metadata_fields:
            prepared.append('')
            prepared.append(_build_metadata_block(metadata_fields, heading_text))
            prepared.append('')
            index = metadata_index

    return '\n'.join(prepared)


def _is_table_row(line):
    s = line.strip()
    return len(s) > 1 and s.startswith('|') and s.endswith('|')


def _is_table_separator(line):
    cells = [c.strip() for c in line.strip().strip('|').split('|')]
    return bool(cells) and all(re.fullmatch(r':?-+:?', c) for c in cells)


def _split_table_row(line):
    return [c.strip() for c in line.strip().strip('|').split('|')]


def _normalize_table_rows(rows):
    """Return (rows, has_header). A leading all-empty row is the markdown
    table's empty header (used for key/value tables); drop it and treat the
    table as headerless so the first column is emphasized instead."""
    if rows and not any(c.strip() for c in rows[0]):
        return rows[1:], False
    return rows, True


def _strip_inline_md(text_value):
    s = re.sub(r'\*\*(.+?)\*\*', r'\1', text_value)
    s = re.sub(r'`([^`]+)`', r'\1', s)
    s = re.sub(r'\[(.+?)\]\((https?://[^\s)]+)\)', r'\1 (\2)', s)
    return s


def _is_horizontal_rule(line):
    return bool(re.fullmatch(r'(-{3,}|\*{3,}|_{3,})', line.strip()))


def _is_block_boundary(line):
    """A line that begins a new block, i.e. cannot be a soft-wrapped continuation
    of the paragraph above it."""
    s = line.strip()
    return (not s) or _is_table_row(s) or _is_horizontal_rule(s) \
        or s.startswith('### ') or s.startswith('## ') or s.startswith('# ') \
        or s.startswith('- ') or s.startswith('> ')


def _docx_add_inline(paragraph, text_value):
    """Add text to a docx paragraph, rendering **bold**, *italic*/_italic_ and
    `code` as runs, decoding HTML entities, honoring hard-break newlines, and
    reducing links to their label (rather than dropping the formatting)."""
    text = html.unescape(text_value)
    text = re.sub(r'\[(.+?)\]\((https?://[^\s)]+)\)', r'\1', text)
    split_pattern = (
        r'(\*\*.+?\*\*|\*(?!\s)[^*\n]+?(?<!\s)\*|(?<![\w])_(?!\s)[^_\n]+?(?<!\s)_(?![\w])|`[^`]+`)'
    )
    lines = text.split('\n')
    for line_index, line in enumerate(lines):
        if line_index:
            paragraph.add_run().add_break()
        for part in re.split(split_pattern, line):
            if not part:
                continue
            if part.startswith('**') and part.endswith('**'):
                paragraph.add_run(part[2:-2]).bold = True
            elif part.startswith('`') and part.endswith('`'):
                paragraph.add_run(part[1:-1])
            elif len(part) >= 2 and part[0] == '*' and part[-1] == '*':
                paragraph.add_run(part[1:-1]).italic = True
            elif len(part) >= 2 and part[0] == '_' and part[-1] == '_':
                paragraph.add_run(part[1:-1]).italic = True
            else:
                paragraph.add_run(part)


def _iter_markdown_blocks(markdown_body):
    lines = markdown_body.replace('\r\n', '\n').split('\n')
    i, n = 0, len(lines)
    while i < n:
        line = lines[i].strip()
        if _is_table_row(line):
            rows = []
            while i < n and _is_table_row(lines[i].strip()):
                current = lines[i].strip()
                if not _is_table_separator(current):
                    rows.append(_split_table_row(current))
                i += 1
            if rows:
                yield ('table', rows)
            continue
        if not line:
            yield ('blank', ''); i += 1; continue
        if _is_horizontal_rule(line):
            yield ('rule', ''); i += 1; continue
        if line.startswith('### '):
            yield ('heading3', line[4:].strip()); i += 1; continue
        if line.startswith('## '):
            yield ('heading2', line[3:].strip()); i += 1; continue
        if line.startswith('# '):
            yield ('heading1', line[2:].strip()); i += 1; continue
        if line.startswith('- '):
            bullet_parts = [line[2:].strip()]
            i += 1
            while i < n and lines[i].strip() and not _is_block_boundary(lines[i]):
                bullet_parts.append(lines[i].strip())
                i += 1
            yield ('bullet', ' '.join(bullet_parts))
            continue
        if line.startswith('> '):
            quote = []
            while i < n and lines[i].strip().startswith('> '):
                quote.append(lines[i].strip()[2:].strip())
                i += 1
            yield ('quote', ' '.join(quote))
            continue
        # paragraph: join soft-wrapped continuation lines into one paragraph,
        # but preserve markdown hard breaks (trailing two spaces / <br>) as newlines
        para_parts = []
        while i < n and not _is_block_boundary(lines[i]):
            raw = lines[i].rstrip('\n')
            hard_break = bool(re.search(r'(  +$|<br\s*/?>\s*$)', raw))
            segment = re.sub(r'\s*<br\s*/?>\s*$', '', raw.strip())
            para_parts.append((segment, hard_break))
            i += 1
        if para_parts:
            buffer = ''
            for part_index, (segment, hard_break) in enumerate(para_parts):
                buffer += segment
                if part_index < len(para_parts) - 1:
                    buffer += '\n' if hard_break else ' '
            yield ('paragraph', buffer)


def _create_export_run(document, version, export_format):
    export_run = ISMSExportRun(
        document_id=document.id,
        document_version_id=version.id,
        export_format=export_format,
        status='completed',
        generated_by=_current_actor_username(),
    )
    db.session.add(export_run)
    _log_action(document, 'isms_export', f'Exported version {version.version_number} as {export_format}')
    db.session.commit()


def _build_markdown_export(document, version):
    output = io.BytesIO(version.markdown_body.encode('utf-8'))
    output.seek(0)
    return send_file(
        output,
        mimetype='text/markdown; charset=utf-8',
        as_attachment=True,
        download_name=_download_filename(document, version.version_number, 'md'),
    )


# ---------------------------------------------------------------------------
# Controlled-document export presentation (cover, TOC, running header/footer).
# Presentation only — never mutates document markdown.
# ---------------------------------------------------------------------------

_LOGO_PATH = Path(__file__).resolve().parent.parent / 'static' / 'images' / 'company_logo.png'
_SLATE = '#222A35'
_SLATE_MUTED = '#5B6370'
_MUTED = '#6B7480'
_RULE = '#C3CBD5'


def _export_document_meta(document, version):
    """Assemble the title-block/header metadata for a controlled document,
    reading whatever fields the document row and its markdown actually carry.
    Missing fields are returned as None and simply omitted downstream."""
    markdown_body = version.markdown_body or ''

    def clean(value):
        value = html.unescape(value)
        value = re.sub(r'\*\*(.+?)\*\*', r'\1', value)
        value = re.sub(r'`([^`]+)`', r'\1', value)
        value = re.sub(r'\s+', ' ', value).strip().strip('·').strip()
        return value or None

    def lookup(*labels):
        for label in labels:
            escaped = re.escape(label)
            inline = re.search(
                r'\*\*\s*' + escaped + r'\s*:?\s*\*\*\s*:?\s*([^\n*|]+)', markdown_body, re.IGNORECASE)
            if inline:
                value = clean(inline.group(1))
                if value:
                    return value
            table = re.search(
                r'\|\s*\*\*\s*' + escaped + r'\s*\*\*\s*\|\s*([^\n|]+?)\s*\|', markdown_body, re.IGNORECASE)
            if table:
                value = clean(table.group(1))
                if value:
                    return value
        return None

    status = lookup('status') or (document.status.title() if document.status else None)
    doc_type = (document.doc_type or '').strip()
    return {
        'code': (document.slug or 'isms-document').upper(),
        'title': document.title or 'Untitled Document',
        'doc_type_label': doc_type.title() if doc_type else None,
        'revision_number': version.version_number,
        'version_label': lookup('version'),
        'status': status,
        'category': document.category or lookup('category'),
        'effective_date': lookup('effective date', 'effective'),
        'classification': lookup('classification'),
        'owner': lookup('owner', 'document owner', 'prepared by', 'prepared'),
        'approver': lookup('approver', 'approved by'),
        'review': lookup('review cycle', 'review date', 'next review'),
        'footer_left': f'Exported from Tracker on {datetime.utcnow().strftime("%Y-%m-%d %H:%M UTC")}',
    }


def _export_meta_rows(meta):
    """Ordered (label, value) pairs for the cover title block, present fields only."""
    version_display = meta['version_label'] or f"v{meta['revision_number']}"
    candidates = [
        ('Document code', meta['code']),
        ('Version', version_display),
        ('Status', meta['status']),
        ('Effective date', meta['effective_date']),
        ('Classification', meta['classification']),
        ('Category', meta['category']),
        ('Owner', meta['owner']),
        ('Approver', meta['approver']),
        ('Review cycle', meta['review']),
        ('Tracker revision', f"v{meta['revision_number']}"),
    ]
    return [(label, value) for label, value in candidates if value]


class _NumberedCanvas(canvas.Canvas):
    """Two-pass canvas so every non-cover page carries a running header
    (doc code / version + rule) and footer (export note / Page X of Y)."""

    def __init__(self, *args, header_left='', header_right='', footer_left='', **kwargs):
        self._header_left = header_left
        self._header_right = header_right
        self._footer_left = footer_left
        super().__init__(*args, **kwargs)
        self._saved_page_states = []

    def showPage(self):
        self._saved_page_states.append(dict(self.__dict__))
        self._startPage()

    def save(self):
        total_pages = len(self._saved_page_states)
        for index, state in enumerate(self._saved_page_states):
            self.__dict__.update(state)
            self._draw_decorations(index + 1, total_pages)
            canvas.Canvas.showPage(self)
        canvas.Canvas.save(self)

    def _draw_decorations(self, page_number, total_pages):
        if page_number == 1:
            return
        width, height = letter
        left = 0.85 * inch
        right = width - 0.85 * inch
        self.saveState()
        self.setFont('Helvetica', 8)
        self.setFillColor(colors.HexColor(_MUTED))
        self.drawString(left, height - 0.6 * inch, self._header_left)
        self.drawRightString(right, height - 0.6 * inch, self._header_right)
        self.setStrokeColor(colors.HexColor(_RULE))
        self.setLineWidth(0.5)
        self.line(left, height - 0.66 * inch, right, height - 0.66 * inch)
        self.line(left, 0.62 * inch, right, 0.62 * inch)
        self.drawString(left, 0.48 * inch, self._footer_left)
        self.drawRightString(right, 0.48 * inch, f'Page {page_number} of {total_pages}')
        self.restoreState()


def _docx_apply_base_styles(doc):
    normal = doc.styles['Normal']
    normal.font.name = 'Cambria'
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.15
    slate = RGBColor(0x22, 0x2A, 0x35)
    for name, size in (('Heading 1', 15), ('Heading 2', 12.5), ('Heading 3', 11)):
        try:
            style = doc.styles[name]
        except KeyError:
            continue
        style.font.name = 'Calibri'
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = slate


def _docx_paragraph_bottom_border(paragraph, color='C3CBD5'):
    p_pr = paragraph._p.get_or_add_pPr()
    borders = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '4')
    bottom.set(qn('w:space'), '2')
    bottom.set(qn('w:color'), color)
    borders.append(bottom)
    p_pr.append(borders)


def _docx_add_page_field(paragraph):
    muted = RGBColor(0x6B, 0x74, 0x80)

    def add_text(text):
        run = paragraph.add_run(text)
        run.font.size = Pt(8)
        run.font.color.rgb = muted

    def add_field(instruction):
        run = paragraph.add_run()
        run.font.size = Pt(8)
        run.font.color.rgb = muted
        begin = OxmlElement('w:fldChar')
        begin.set(qn('w:fldCharType'), 'begin')
        instr = OxmlElement('w:instrText')
        instr.set(qn('xml:space'), 'preserve')
        instr.text = f' {instruction} '
        end = OxmlElement('w:fldChar')
        end.set(qn('w:fldCharType'), 'end')
        run._r.append(begin)
        run._r.append(instr)
        run._r.append(end)

    add_text('Page ')
    add_field('PAGE')
    add_text(' of ')
    add_field('NUMPAGES')


def _docx_setup_running_headers(doc, meta):
    section = doc.sections[0]
    section.different_first_page_header_footer = True
    content_width = section.page_width - section.left_margin - section.right_margin
    muted = RGBColor(0x6B, 0x74, 0x80)

    header = section.header
    header.is_linked_to_previous = False
    header_paragraph = header.paragraphs[0]
    header_paragraph.text = ''
    header_paragraph.paragraph_format.tab_stops.add_tab_stop(content_width, WD_TAB_ALIGNMENT.RIGHT)
    left_run = header_paragraph.add_run(meta['code'])
    left_run.font.size = Pt(8)
    left_run.font.color.rgb = muted
    header_paragraph.add_run('\t')
    right_run = header_paragraph.add_run(f"v{meta['revision_number']}")
    right_run.font.size = Pt(8)
    right_run.font.color.rgb = muted
    _docx_paragraph_bottom_border(header_paragraph)

    footer = section.footer
    footer.is_linked_to_previous = False
    footer_paragraph = footer.paragraphs[0]
    footer_paragraph.text = ''
    footer_paragraph.paragraph_format.tab_stops.add_tab_stop(content_width, WD_TAB_ALIGNMENT.RIGHT)
    note_run = footer_paragraph.add_run(meta['footer_left'])
    note_run.font.size = Pt(8)
    note_run.font.color.rgb = muted
    footer_paragraph.add_run('\t')
    _docx_add_page_field(footer_paragraph)


def _docx_build_cover(doc, meta):
    if _LOGO_PATH.exists():
        logo_paragraph = doc.add_paragraph()
        logo_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        logo_paragraph.add_run().add_picture(str(_LOGO_PATH), width=Inches(0.95))

    doc.add_paragraph()
    doc.add_paragraph()

    code_paragraph = doc.add_paragraph()
    code_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    code_run = code_paragraph.add_run(meta['code'])
    code_run.bold = True
    code_run.font.name = 'Calibri'
    code_run.font.size = Pt(11)
    code_run.font.color.rgb = RGBColor(0x5B, 0x63, 0x70)

    title_paragraph = doc.add_paragraph()
    title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_paragraph.add_run(meta['title'])
    title_run.bold = True
    title_run.font.name = 'Calibri'
    title_run.font.size = Pt(24)
    title_run.font.color.rgb = RGBColor(0x22, 0x2A, 0x35)

    if meta['doc_type_label']:
        subtitle_paragraph = doc.add_paragraph()
        subtitle_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        subtitle_run = subtitle_paragraph.add_run(meta['doc_type_label'])
        subtitle_run.italic = True
        subtitle_run.font.size = Pt(11)
        subtitle_run.font.color.rgb = RGBColor(0x6B, 0x74, 0x80)

    doc.add_paragraph()
    doc.add_paragraph()

    rows = _export_meta_rows(meta)
    if rows:
        table = doc.add_table(rows=len(rows), cols=2)
        for row_index, (label, value) in enumerate(rows):
            label_cell = table.rows[row_index].cells[0]
            value_cell = table.rows[row_index].cells[1]
            label_cell.width = Inches(1.9)
            value_cell.width = Inches(3.9)
            label_run = label_cell.paragraphs[0].add_run(label)
            label_run.bold = True
            label_run.font.size = Pt(10)
            label_run.font.color.rgb = RGBColor(0x39, 0x43, 0x4F)
            # value is already plain text (entities decoded, markdown stripped) from
            # _export_document_meta.clean(); add it as a single run to avoid a second decode.
            value_run = value_cell.paragraphs[0].add_run(value)
            value_run.font.size = Pt(10)

    doc.add_page_break()


def _docx_build_toc(doc, version):
    toc_items = render_markdown_with_toc(version.markdown_body).get('toc') or []
    if not toc_items:
        return
    heading = doc.add_paragraph()
    heading_run = heading.add_run('Contents')
    heading_run.bold = True
    heading_run.font.name = 'Calibri'
    heading_run.font.size = Pt(15)
    heading_run.font.color.rgb = RGBColor(0x22, 0x2A, 0x35)
    heading.paragraph_format.space_after = Pt(8)
    for item in toc_items:
        level = item.get('level', 1)
        entry = doc.add_paragraph()
        entry.paragraph_format.left_indent = Inches(0.28 * (level - 1))
        entry.paragraph_format.space_after = Pt(2)
        entry_run = entry.add_run(item.get('title', ''))
        if level == 1:
            entry_run.bold = True
            entry_run.font.size = Pt(11)
        else:
            entry_run.font.size = Pt(10)
            entry_run.font.color.rgb = RGBColor(0x39, 0x43, 0x4F)
    doc.add_page_break()


def _build_docx_export(document, version):
    meta = _export_document_meta(document, version)
    doc = Document()
    _docx_apply_base_styles(doc)
    _docx_setup_running_headers(doc, meta)
    _docx_build_cover(doc, meta)
    _docx_build_toc(doc, version)

    for block_type, text_value in _iter_markdown_blocks(version.markdown_body):
        if block_type == 'blank':
            continue
        if block_type == 'heading1':
            doc.add_heading(text_value, level=1)
        elif block_type == 'heading2':
            doc.add_heading(text_value, level=2)
        elif block_type == 'heading3':
            doc.add_heading(text_value, level=3)
        elif block_type == 'bullet':
            _docx_add_inline(doc.add_paragraph(style='List Bullet'), text_value)
        elif block_type == 'quote':
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.3)
            _docx_add_inline(p, text_value)
            for run in p.runs:
                run.italic = True
        elif block_type == 'rule':
            continue
        elif block_type == 'table':
            rows, has_header = _normalize_table_rows(text_value)
            if not rows:
                continue
            ncols = max(len(r) for r in rows)
            table = doc.add_table(rows=len(rows), cols=ncols)
            table.style = 'Table Grid'
            for ri, row in enumerate(rows):
                for ci in range(ncols):
                    cell = table.rows[ri].cells[ci]
                    cell.text = _strip_inline_md(row[ci]) if ci < len(row) else ''
                    emphasize = (ri == 0 and has_header) or (not has_header and ci == 0)
                    for para in cell.paragraphs:
                        for run in para.runs:
                            run.font.size = Pt(9)
                            if emphasize:
                                run.bold = True
            doc.add_paragraph()
        else:
            _docx_add_inline(doc.add_paragraph(), text_value)

    output = io.BytesIO()
    doc.save(output)
    output.seek(0)
    return send_file(
        output,
        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        as_attachment=True,
        download_name=_download_filename(document, version.version_number, 'docx'),
    )


def _build_pdf_export(document, version):
    meta = _export_document_meta(document, version)
    output = io.BytesIO()
    content_width = letter[0] - (2 * 0.85 * inch)
    pdf = SimpleDocTemplate(
        output,
        pagesize=letter,
        topMargin=0.95 * inch,
        bottomMargin=0.85 * inch,
        leftMargin=0.85 * inch,
        rightMargin=0.85 * inch,
        title=meta['title'],
        author='Cirque Corporation',
        subject=meta['code'],
    )
    styles = getSampleStyleSheet()
    heading1_style = ParagraphStyle(
        'ISMSExportHeading1', parent=styles['Heading2'], fontName='Helvetica-Bold',
        fontSize=14, leading=17, textColor=colors.HexColor(_SLATE), spaceBefore=12, spaceAfter=7,
    )
    heading2_style = ParagraphStyle(
        'ISMSExportHeading2', parent=styles['Heading3'], fontName='Helvetica-Bold',
        fontSize=12, leading=15, textColor=colors.HexColor(_SLATE), spaceBefore=10, spaceAfter=5,
    )
    heading3_style = ParagraphStyle(
        'ISMSExportHeading3', parent=styles['Heading4'], fontName='Helvetica-Bold',
        fontSize=10.5, leading=13, textColor=colors.HexColor('#39434F'), spaceBefore=8, spaceAfter=4,
    )
    body_style = ParagraphStyle(
        'ISMSExportBody', parent=styles['BodyText'], fontName='Times-Roman',
        fontSize=10.5, leading=14.5, spaceAfter=8,
    )
    bullet_style = ParagraphStyle(
        'ISMSExportBullet', parent=body_style, leftIndent=18, firstLineIndent=-8, spaceAfter=5,
    )
    table_cell_style = ParagraphStyle(
        'ISMSExportTableCell', parent=body_style, fontSize=8.5, leading=11, spaceAfter=0,
    )
    table_header_style = ParagraphStyle(
        'ISMSExportTableHeader', parent=table_cell_style, fontName='Times-Bold',
    )
    quote_style = ParagraphStyle(
        'ISMSExportQuote', parent=body_style, leftIndent=16, fontName='Times-Italic',
        textColor=colors.HexColor('#39434F'), spaceBefore=4, spaceAfter=8,
    )
    cover_code_style = ParagraphStyle(
        'ISMSCoverCode', parent=body_style, fontName='Helvetica-Bold', fontSize=11,
        textColor=colors.HexColor(_SLATE_MUTED), alignment=TA_CENTER, spaceAfter=6,
    )
    cover_title_style = ParagraphStyle(
        'ISMSCoverTitle', parent=body_style, fontName='Helvetica-Bold', fontSize=24, leading=28,
        textColor=colors.HexColor(_SLATE), alignment=TA_CENTER, spaceAfter=4,
    )
    cover_subtitle_style = ParagraphStyle(
        'ISMSCoverSubtitle', parent=body_style, fontName='Helvetica-Oblique', fontSize=11,
        textColor=colors.HexColor(_MUTED), alignment=TA_CENTER, spaceAfter=4,
    )
    cover_label_style = ParagraphStyle(
        'ISMSCoverLabel', parent=body_style, fontName='Helvetica-Bold', fontSize=9.5, leading=13,
        textColor=colors.HexColor('#39434F'), spaceAfter=0,
    )
    cover_value_style = ParagraphStyle(
        'ISMSCoverValue', parent=body_style, fontSize=9.5, leading=13, spaceAfter=0,
    )
    toc_title_style = ParagraphStyle(
        'ISMSTocTitle', parent=heading1_style, fontSize=15, spaceBefore=0, spaceAfter=8,
    )
    toc_level_styles = {
        1: ParagraphStyle('ISMSToc1', parent=body_style, fontName='Helvetica-Bold', fontSize=10.5,
                          leading=15, spaceAfter=1),
        2: ParagraphStyle('ISMSToc2', parent=body_style, fontSize=10, leading=14, leftIndent=16,
                          textColor=colors.HexColor('#39434F'), spaceAfter=1),
        3: ParagraphStyle('ISMSToc3', parent=body_style, fontSize=9.5, leading=13, leftIndent=32,
                          textColor=colors.HexColor(_SLATE_MUTED), spaceAfter=1),
    }

    # --- Cover page ---
    story = []
    if _LOGO_PATH.exists():
        logo = Image(str(_LOGO_PATH), width=0.95 * inch, height=0.95 * inch * (174.0 / 187.0))
        logo.hAlign = 'CENTER'
        story.append(logo)
        story.append(Spacer(1, 0.35 * inch))
    else:
        story.append(Spacer(1, 1.2 * inch))
    story.append(Paragraph(html.escape(meta['code']), cover_code_style))
    story.append(Paragraph(html.escape(meta['title']), cover_title_style))
    if meta['doc_type_label']:
        story.append(Paragraph(html.escape(meta['doc_type_label']), cover_subtitle_style))
    story.append(Spacer(1, 0.4 * inch))
    meta_rows = _export_meta_rows(meta)
    if meta_rows:
        data = [[Paragraph(html.escape(label), cover_label_style),
                 Paragraph(html.escape(value), cover_value_style)] for label, value in meta_rows]
        meta_table = Table(data, colWidths=[1.7 * inch, 3.9 * inch], hAlign='CENTER')
        meta_table.setStyle(TableStyle([
            ('VALIGN', (0, 0), (-1, -1), 'TOP'),
            ('LEFTPADDING', (0, 0), (-1, -1), 0),
            ('RIGHTPADDING', (0, 0), (-1, -1), 4),
            ('TOPPADDING', (0, 0), (-1, -1), 4),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 4),
            ('LINEBELOW', (0, 0), (-1, -2), 0.4, colors.HexColor('#E2E7ED')),
        ]))
        story.append(meta_table)
    story.append(PageBreak())

    # --- Table of contents ---
    toc_items = render_markdown_with_toc(version.markdown_body).get('toc') or []
    if toc_items:
        story.append(Paragraph('Contents', toc_title_style))
        for item in toc_items:
            level = min(item.get('level', 1), 3)
            story.append(Paragraph(_render_inline_export(item.get('title', '')), toc_level_styles[level]))
        story.append(PageBreak())

    # --- Body ---
    for block_type, text_value in _iter_markdown_blocks(version.markdown_body):
        if block_type == 'blank':
            story.append(Spacer(1, 0.06 * inch))
        elif block_type == 'heading1':
            story.append(Paragraph(_render_inline_export(text_value), heading1_style))
        elif block_type == 'heading2':
            story.append(Paragraph(_render_inline_export(text_value), heading2_style))
        elif block_type == 'heading3':
            story.append(Paragraph(_render_inline_export(text_value), heading3_style))
        elif block_type == 'bullet':
            story.append(Paragraph(f'&bull;&nbsp; {_render_inline_export(text_value)}', bullet_style))
        elif block_type == 'quote':
            story.append(Paragraph(_render_inline_export(text_value), quote_style))
        elif block_type == 'rule':
            story.append(HRFlowable(width='100%', thickness=0.5,
                                    color=colors.HexColor(_RULE), spaceBefore=6, spaceAfter=8))
        elif block_type == 'table':
            rows, has_header = _normalize_table_rows(text_value)
            if rows:
                ncols = max(len(r) for r in rows)
                col_w = content_width / ncols
                data = []
                for row in rows:
                    row_style = table_header_style if (has_header and not data) else table_cell_style
                    cells = [Paragraph(_render_inline_export(row[ci] if ci < len(row) else ''), row_style)
                             for ci in range(ncols)]
                    data.append(cells)
                tbl = Table(data, colWidths=[col_w] * ncols, repeatRows=1 if has_header else 0)
                tstyle = [
                    ('GRID', (0, 0), (-1, -1), 0.5, colors.HexColor(_RULE)),
                    ('VALIGN', (0, 0), (-1, -1), 'TOP'),
                    ('LEFTPADDING', (0, 0), (-1, -1), 5),
                    ('RIGHTPADDING', (0, 0), (-1, -1), 5),
                    ('TOPPADDING', (0, 0), (-1, -1), 4),
                    ('BOTTOMPADDING', (0, 0), (-1, -1), 4),
                ]
                if has_header:
                    tstyle.append(('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#EEF2F6')))
                    tstyle.append(('ROWBACKGROUNDS', (0, 1), (-1, -1),
                                   [colors.white, colors.HexColor('#F7F9FB')]))
                else:
                    tstyle.append(('BACKGROUND', (0, 0), (0, -1), colors.HexColor('#F6F8FA')))
                tbl.setStyle(TableStyle(tstyle))
                story.append(tbl)
                story.append(Spacer(1, 0.12 * inch))
        else:
            story.append(Paragraph(_render_inline_export(text_value), body_style))

    pdf.build(story, canvasmaker=partial(
        _NumberedCanvas,
        header_left=meta['code'],
        header_right=f"v{meta['revision_number']}",
        footer_left=meta['footer_left'],
    ))
    output.seek(0)
    return send_file(
        output,
        mimetype='application/pdf',
        as_attachment=True,
        download_name=_download_filename(document, version.version_number, 'pdf'),
    )


def _parse_front_matter(markdown_body):
    if not markdown_body.startswith('---\n'):
        return {}

    parts = markdown_body.split('\n---\n', 1)
    if len(parts) != 2:
        return {}

    metadata = {}
    for line in parts[0].splitlines()[1:]:
        if ':' not in line:
            continue
        key, value = line.split(':', 1)
        metadata[key.strip().lower()] = value.strip().strip('"').strip("'")
    return metadata


def _parse_manual_version_label(markdown_body):
    match = re.search(r'^\|\s*([0-9]+(?:\.[0-9]+)?)\s*\|\s*\d{4}-\d{2}-\d{2}\s*\|', markdown_body, re.MULTILINE)
    return match.group(1) if match else '1.0'


def _normalize_retention_status(retention_value):
    normalized = ' '.join(str(retention_value or '').strip().lower().split())
    status_map = {
        'exist and no corrections': 'Exist & No corrections',
        'exist & no corrections': 'Exist & No corrections',
        'exist and corrections': 'Exist & corrections',
        'exist & corrections': 'Exist & corrections',
    }
    return status_map.get(normalized)


def _normalize_effective_date(date_value):
    raw_value = str(date_value or '').strip()
    if not raw_value:
        return None

    for fmt in ('%Y-%m-%d', '%B %Y', '%b %Y'):
        try:
            parsed = datetime.strptime(raw_value, fmt)
            if fmt in ('%B %Y', '%b %Y'):
                parsed = parsed.replace(day=1)
            return parsed.strftime('%Y-%m-%d')
        except ValueError:
            continue

    return raw_value


def _parse_isms_manual_documents(markdown_body):
    lines = markdown_body.replace('\r\n', '\n').split('\n')
    documents = {}

    for index, raw_line in enumerate(lines):
        heading_match = re.match(r'^##\s+(IS-[A-Z0-9-]+):\s*(.+)$', raw_line.strip())
        if not heading_match:
            continue

        heading_document_id, heading_title = heading_match.groups()
        metadata = {}
        cursor = index + 1
        while cursor < len(lines):
            candidate = lines[cursor].strip()
            if candidate.startswith('## '):
                break
            parsed_fields = _extract_metadata_fields(candidate)
            if parsed_fields:
                for label, value in parsed_fields:
                    metadata[label.strip()] = value.strip().strip('*').strip()
            cursor += 1

        document_id = metadata.get('Document ID') or metadata.get('Document') or heading_document_id
        documents[document_id] = {
            'document_id': document_id,
            'title': metadata.get('Document Title') or metadata.get('Standards Name') or heading_title.strip(),
            'effective_date': _normalize_effective_date(metadata.get('Effective Date')),
            'category': metadata.get('Category'),
            'division': metadata.get('Division'),
            'document_type': metadata.get('Document Type') or metadata.get('Standard Type'),
            'version': metadata.get('Version'),
            'retention_status': _normalize_retention_status(metadata.get('Standard Retention')),
        }

    return documents


def _extract_document_ids(cell_value):
    document_ids = []
    last_prefix = None

    for raw_token in re.split(r'\s*\+\s*', str(cell_value or '')):
        token = raw_token.strip()
        if not token:
            continue

        full_match = re.search(r'IS-[A-Z0-9-]+', token)
        if full_match:
            document_id = full_match.group(0)
            document_ids.append(document_id)
            prefix_match = re.match(r'^(IS-[A-Z0-9]+-)', document_id)
            if prefix_match:
                last_prefix = prefix_match.group(1)
            continue

        shorthand_match = re.match(r'^(CIRQ[0-9]{2}-[A-Z0-9]+)$', token)
        if shorthand_match and last_prefix:
            document_ids.append(f'{last_prefix}{shorthand_match.group(1)}')

    return document_ids


def _refresh_inventory_sheet(worksheet, manual_documents, manual_date, manual_version_label):
    worksheet['A2'] = (
        f'Source: ISMS-Manual.docx v{manual_version_label} '
        f'(effective {manual_date}, A4, Arial 10pt, parent-compliant format)'
    )

    for row in range(6, worksheet.max_row + 1):
        cirque_document_id = worksheet.cell(row=row, column=2).value
        if not cirque_document_id:
            continue

        manual_document = manual_documents.get(str(cirque_document_id).strip())
        if manual_document is None:
            continue

        if manual_document.get('title'):
            worksheet.cell(row=row, column=6, value=manual_document['title'])
        if manual_document.get('retention_status'):
            worksheet.cell(row=row, column=7, value=manual_document['retention_status'])
        if manual_document.get('effective_date'):
            worksheet.cell(row=row, column=9, value=manual_document['effective_date'])


def _refresh_survey_sheet(worksheet, manual_documents, tracked_asset_count):
    total_rows = 0
    correction_rows = 0

    for row in range(3, worksheet.max_row + 1):
        reference_number = worksheet.cell(row=row, column=1).value
        if reference_number in (None, ''):
            continue

        total_rows += 1
        mapped_document_ids = _extract_document_ids(worksheet.cell(row=row, column=10).value)
        mapped_documents = [manual_documents[document_id] for document_id in mapped_document_ids if document_id in manual_documents]
        effective_dates = [item['effective_date'] for item in mapped_documents if item.get('effective_date')]
        if effective_dates:
            worksheet.cell(row=row, column=11, value=max(effective_dates))

        mapped_statuses = {item['retention_status'] for item in mapped_documents if item.get('retention_status')}
        if mapped_statuses and mapped_statuses == {'Exist & No corrections'}:
            worksheet.cell(row=row, column=7, value='Exist & No corrections')

        if tracked_asset_count and worksheet.cell(row=row, column=10).value and 'IS-AAR01-CIRQ01-F01A' in str(worksheet.cell(row=row, column=10).value):
            current_reason = str(worksheet.cell(row=row, column=8).value or '').strip()
            asset_suffix = f' Tracker currently maintains {tracked_asset_count} assets.'
            if current_reason and asset_suffix not in current_reason:
                worksheet.cell(row=row, column=8, value=current_reason.rstrip('.') + '.' + asset_suffix)

        override = GLOCALIZATION_SURVEY_OVERRIDES.get(reference_number)
        if override:
            if override.get('status'):
                worksheet.cell(row=row, column=7, value=override['status'])
            if override.get('reason'):
                worksheet.cell(row=row, column=8, value=override['reason'])
            if override.get('plan_date'):
                worksheet.cell(row=row, column=11, value=override['plan_date'])

        status = str(worksheet.cell(row=row, column=7).value or '').strip()
        if status == 'Exist & corrections':
            correction_rows += 1

    return total_rows, correction_rows


def _refresh_progress_sheet(worksheet, total_rows, correction_rows):
    completed_rows = max(total_rows - correction_rows, 0)
    worksheet['C5'] = 1
    worksheet['D5'] = 1
    worksheet['E5'] = '=D5/C5'
    worksheet['F5'] = 'Filled and ready; generated from the current Tracker workbook template'

    worksheet['C6'] = 0
    worksheet['D6'] = 0
    worksheet['E6'] = 'N/A'
    worksheet['F6'] = 'Cirque has no Global Version adoptions; not applicable'

    worksheet['C7'] = total_rows
    worksheet['D7'] = completed_rows
    worksheet['E7'] = '=D7/C7' if total_rows else 'N/A'
    worksheet['F7'] = (
        f'{completed_rows}/{total_rows} ready as-is; {correction_rows} pending edits'
        if correction_rows else
        'All localized/local rows align with the current manual metadata'
    )

    worksheet['C8'] = total_rows
    worksheet['D8'] = completed_rows
    worksheet['E8'] = '=D8/C8' if total_rows else 'N/A'
    worksheet['F8'] = 'Approval tracked in Cirque Inventory'

    worksheet['C9'] = total_rows
    worksheet['D9'] = 0
    worksheet['E9'] = '=D9/C9' if total_rows else 'N/A'
    worksheet['F9'] = 'None uploaded yet'


def _build_glocalization_survey_export(document, version):
    if not GLOCALIZATION_SURVEY_TEMPLATE_PATH.exists():
        raise FileNotFoundError(f'Glocalization survey template not found: {GLOCALIZATION_SURVEY_TEMPLATE_PATH}')

    workbook = load_workbook(GLOCALIZATION_SURVEY_TEMPLATE_PATH)
    manual_front_matter = _parse_front_matter(version.markdown_body)
    manual_documents = _parse_isms_manual_documents(version.markdown_body)
    manual_date = manual_front_matter.get('date', datetime.utcnow().strftime('%Y-%m-%d'))
    manual_version_label = _parse_manual_version_label(version.markdown_body)
    tracked_asset_count = Asset.query.count()

    progress_sheet = workbook['Progress']
    inventory_sheet = workbook['Cirque Inventory']
    survey_sheet = workbook['Survey Sheet']

    _refresh_inventory_sheet(inventory_sheet, manual_documents, manual_date, manual_version_label)
    total_rows, correction_rows = _refresh_survey_sheet(survey_sheet, manual_documents, tracked_asset_count)
    _refresh_progress_sheet(progress_sheet, total_rows, correction_rows)

    output = io.BytesIO()
    workbook.save(output)
    output.seek(0)
    return output


def _build_diff_table(from_version, to_version):
    diff = HtmlDiff(wrapcolumn=100)
    return diff.make_table(
        from_version.markdown_body.splitlines(),
        to_version.markdown_body.splitlines(),
        fromdesc=f'v{from_version.version_number}',
        todesc=f'v{to_version.version_number}',
        context=True,
        numlines=3,
    )


def _build_diff_summary(from_version, to_version):
    added = 0
    removed = 0
    changed = 0
    matcher = SequenceMatcher(
        a=from_version.markdown_body.splitlines(),
        b=to_version.markdown_body.splitlines(),
    )

    for tag, a_start, a_end, b_start, b_end in matcher.get_opcodes():
        if tag == 'insert':
            added += b_end - b_start
        elif tag == 'delete':
            removed += a_end - a_start
        elif tag == 'replace':
            changed += max(a_end - a_start, b_end - b_start)

    return {
        'added': added,
        'removed': removed,
        'changed': changed,
    }


def _resolve_diff_versions(document, selected_version_id=None, target_version_id=None):
    versions = document.versions
    current_version = _get_current_version(document)
    if current_version is None:
        abort(404)

    version_by_id = {version.id: version for version in versions}
    compare_version = version_by_id.get(selected_version_id) if selected_version_id else None
    if compare_version is None:
        compare_version = next((version for version in versions if version.id != current_version.id), None)
    target_version = version_by_id.get(target_version_id) if target_version_id else current_version

    if compare_version is None or target_version is None:
        abort(404)
    if compare_version.document_id != document.id or target_version.document_id != document.id:
        abort(404)
    if compare_version.id == target_version.id:
        flash('Choose two different versions to compare.', 'info')
        return None, None, versions

    return compare_version, target_version, versions


def _render_inline(text_value):
    """Reader/metadata-card inline renderer. MUST stay byte-identical to its
    historical behavior — the on-screen isms-section-meta-value blocks depend on
    it. Handles **bold** and links only (no italics, no entity decoding). The
    export path uses _render_inline_export instead."""
    rendered = html.escape(text_value)
    rendered = re.sub(r'\*\*(.+?)\*\*', r'<strong>\1</strong>', rendered)
    rendered = re.sub(r'\[(.+?)\]\((https?://[^\s)]+)\)', r'<a href="\2" target="_blank" rel="noopener noreferrer">\1</a>', rendered)
    return rendered


def _emphasize_export(fragment):
    """Apply **bold** and *italic*/_italic_ to a non-link text fragment."""
    fragment = re.sub(r'\*\*(.+?)\*\*', r'<strong>\1</strong>', fragment)
    fragment = re.sub(r'(?<![\w*])\*(?!\s)([^*\n]+?)(?<!\s)\*(?![\w*])', r'<em>\1</em>', fragment)
    fragment = re.sub(r'(?<![\w_])_(?!\s)([^_\n]+?)(?<!\s)_(?![\w_])', r'<em>\1</em>', fragment)
    return fragment


def _render_inline_export(text_value):
    """Inline renderer for DOCX/PDF export only: decodes HTML entities, renders
    **bold**, *italic*/_italic_ and links, and preserves hard-break newlines.
    Emphasis is applied only OUTSIDE link href spans, so URLs containing
    underscores (e.g. /_a_/) are never corrupted with <em> tags."""
    text = html.escape(html.unescape(text_value))
    link_re = re.compile(r'\[(.+?)\]\((https?://[^\s)]+)\)')
    parts = []
    cursor = 0
    for match in link_re.finditer(text):
        parts.append(_emphasize_export(text[cursor:match.start()]))
        label = _emphasize_export(match.group(1))
        href = match.group(2)
        # reportlab's Paragraph parser only accepts href/color (not rel/target),
        # so emit a minimal link tag; the href stays intact including underscores.
        parts.append(f'<a href="{href}" color="#1A4FC4">{label}</a>')
        cursor = match.end()
    parts.append(_emphasize_export(text[cursor:]))
    return ''.join(parts).replace('\n', '<br/>')


def render_markdown_with_toc(markdown_body):
    if not markdown_body:
        return {
            'html': '<p class="text-muted">No content available.</p>',
            'toc': [],
        }
    normalized_markdown = _prepare_reader_markdown(markdown_body)
    md = markdown_lib.Markdown(
        extensions=['extra', 'sane_lists', 'toc', 'tables', 'fenced_code'],
        extension_configs={
            'toc': {
                'permalink': False,
            }
        },
        output_format='html5',
    )
    html_output = md.convert(normalized_markdown)
    html_output = re.sub(r'<h([1-6]) id="([^"]+)"', r'<h\1 id="\2" class="isms-heading-anchor"', html_output)
    return {
        'html': html_output,
        'toc': _flatten_toc_tokens(getattr(md, 'toc_tokens', [])),
    }


def render_markdown(markdown_body):
    return render_markdown_with_toc(markdown_body)['html']


@bp.route('/isms')
@login_required
@admin_required
def documents():
    documents = ISMSDocument.query.order_by(ISMSDocument.title.asc()).all()
    manual_document = next((document for document in documents if document.slug == 'isms-manual'), None)
    return render_template('isms_documents.html', documents=documents, manual_document=manual_document)


@bp.route('/isms/<int:document_id>')
@login_required
@admin_required
def document_detail(document_id):
    document = ISMSDocument.query.get_or_404(document_id)
    version = _get_current_version(document)
    if version is None:
        abort(404)

    rendered_document = render_markdown_with_toc(version.markdown_body)
    return render_template(
        'isms_document_detail.html',
        document=document,
        version=version,
        rendered_html=rendered_document['html'],
        toc_items=rendered_document['toc'],
        current_actor=_current_actor(),
    )


@bp.route('/isms/<int:document_id>/edit', methods=['GET', 'POST'])
@login_required
@admin_required
def edit_document(document_id):
    document = ISMSDocument.query.get_or_404(document_id)
    version = _get_current_version(document)
    if version is None:
        abort(404)

    if request.method == 'POST':
        base_version_id = request.form.get('base_version_id', type=int)
        markdown_body = request.form.get('markdown_body', '')
        change_summary = (request.form.get('change_summary') or '').strip()

        if not change_summary:
            flash('Change summary is required.', 'warning')
            return render_template('isms_edit_document.html', document=document, version=version, rendered_html=render_markdown(markdown_body))

        if base_version_id != document.current_version_id:
            flash('This document changed after you opened it. Reload the editor and try again.', 'danger')
            return redirect(url_for('isms.edit_document', document_id=document.id))

        if markdown_body == version.markdown_body:
            flash('No content changes detected. Nothing was saved.', 'info')
            return redirect(url_for('isms.document_detail', document_id=document.id))

        new_version = _create_version(document, markdown_body, change_summary)
        _log_action(document, 'isms_edit', f'Saved version {new_version.version_number}: {change_summary}')
        db.session.commit()
        flash(f'Saved {document.title} as version {new_version.version_number}.', 'success')
        return redirect(url_for('isms.document_detail', document_id=document.id))

    return render_template('isms_edit_document.html', document=document, version=version, rendered_html=version.rendered_html or render_markdown(version.markdown_body))


@bp.route('/isms/<int:document_id>/history')
@login_required
@admin_required
def document_history(document_id):
    document = ISMSDocument.query.get_or_404(document_id)
    version = _get_current_version(document)
    if version is None:
        abort(404)
    return render_template('isms_document_history.html', document=document, current_version=version, versions=document.versions)


@bp.route('/isms/<int:document_id>/diff')
@bp.route('/isms/<int:document_id>/diff/<int:version_id>')
@login_required
@admin_required
def document_diff(document_id, version_id=None):
    document = ISMSDocument.query.get_or_404(document_id)
    selected_version_id = version_id or request.args.get('from_version_id', type=int)
    target_version_id = request.args.get('to_version_id', type=int)

    compare_version, target_version, versions = _resolve_diff_versions(document, selected_version_id, target_version_id)
    if compare_version is None or target_version is None:
        return redirect(url_for('isms.document_history', document_id=document.id))

    diff_table = _build_diff_table(compare_version, target_version)
    diff_summary = _build_diff_summary(compare_version, target_version)
    return render_template(
        'isms_document_diff.html',
        document=document,
        current_version=_get_current_version(document),
        compare_version=compare_version,
        target_version=target_version,
        versions=versions,
        diff_summary=diff_summary,
        diff_table=diff_table,
    )


@bp.route('/isms/<int:document_id>/restore/<int:version_id>', methods=['POST'])
@login_required
@admin_required
def restore_document_version(document_id, version_id):
    document = ISMSDocument.query.get_or_404(document_id)
    restore_source = ISMSDocumentVersion.query.filter_by(document_id=document.id, id=version_id).first_or_404()
    current_version = _get_current_version(document)
    if current_version and current_version.id == restore_source.id:
        flash('That version is already current.', 'info')
        return redirect(url_for('isms.document_history', document_id=document.id))

    restored_version = _create_version(
        document,
        restore_source.markdown_body,
        f'Restored from version {restore_source.version_number}',
        is_restore=True,
        restored_from=restore_source,
    )
    _log_action(document, 'isms_restore', f'Restored version {restore_source.version_number} into new version {restored_version.version_number}')
    db.session.commit()
    flash(f'Restored version {restore_source.version_number} as new current version {restored_version.version_number}.', 'success')
    return redirect(url_for('isms.document_detail', document_id=document.id))


@bp.route('/isms/<int:document_id>/export/<string:export_format>')
@login_required
@admin_required
def export_document(document_id, export_format):
    document = ISMSDocument.query.get_or_404(document_id)
    version = _get_current_version(document)
    if version is None:
        abort(404)

    export_format = (export_format or '').lower()
    if export_format not in {'md', 'docx', 'pdf'}:
        abort(404)

    _create_export_run(document, version, export_format)

    if export_format == 'md':
        return _build_markdown_export(document, version)
    if export_format == 'docx':
        return _build_docx_export(document, version)
    return _build_pdf_export(document, version)


@bp.route('/isms/<int:document_id>/export-glocalization-survey')
@login_required
@admin_required
def export_glocalization_survey(document_id):
    document = ISMSDocument.query.get_or_404(document_id)
    if document.slug != 'isms-manual':
        abort(404)

    version = _get_current_version(document)
    if version is None:
        abort(404)

    _create_export_run(document, version, 'glocalization-xlsx')
    _log_action(document, 'isms_export', f'Exported glocalization survey workbook from version {version.version_number}')

    output = _build_glocalization_survey_export(document, version)
    export_timestamp = datetime.utcnow().strftime('%Y-%m-%d-%H%M%S')
    download_name = f'Glocalization-Survey-Cirque-Filled-V7-{export_timestamp}.xlsx'
    response = send_file(
        output,
        mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
        as_attachment=True,
        download_name=download_name,
    )
    response.headers['Cache-Control'] = 'no-store, no-cache, must-revalidate, max-age=0'
    response.headers['Pragma'] = 'no-cache'
    response.headers['Expires'] = '0'
    return response

# ---------------------------------------------------------------------------
# ISMS Management Ledgers (ALAP IS-APM02-F02..F09 workbook)
# ---------------------------------------------------------------------------

def _log_ledger_action(action, details):
    try:
        db.session.add(
            AuditTrail(
                asset_id=None,
                action=action,
                table_name='isms_ledger',
                record_id=None,
                old_values=None,
                new_values=details,
                changed_by=_current_actor_username(),
            )
        )
        db.session.commit()
    except Exception:
        db.session.rollback()


@bp.route('/isms/ledgers')
@login_required
@admin_required
def ledgers():
    template = isms_ledger_service.template_info()
    coverage = isms_ledger_service.ledger_coverage() if template else None
    readiness = None
    if template and request.args.get('readiness') == '1':
        # Generating the workbook to measure it is not free, so it's opt-in.
        try:
            readiness = isms_ledger_service.ledger_readiness()
        except Exception as exc:
            flash(f'Could not compute readiness: {exc}', 'warning')
    return render_template(
        'isms_ledgers.html',
        template=template,
        coverage=coverage,
        template_dir=str(isms_ledger_service.LEDGER_TEMPLATE_DIR),
        exclusions=isms_ledger_service.ledger_exclusions(),
        exclusion_setting=isms_ledger_service.LEDGER_EXCLUSION_SETTING,
        readiness=readiness,
    )


@bp.route('/isms/ledgers/template', methods=['POST'])
@login_required
@admin_required
def upload_ledger_template():
    upload = request.files.get('template')
    if upload is None or not upload.filename:
        flash('Choose a .xlsx ledger template to upload.', 'warning')
        return redirect(url_for('isms.ledgers'))

    filename = secure_filename(upload.filename)
    if not filename.lower().endswith('.xlsx'):
        flash('The ledger template must be an .xlsx workbook.', 'danger')
        return redirect(url_for('isms.ledgers'))

    isms_ledger_service.LEDGER_TEMPLATE_DIR.mkdir(parents=True, exist_ok=True)
    destination = isms_ledger_service.LEDGER_TEMPLATE_DIR / filename
    upload.save(destination)

    try:
        load_workbook(destination)
    except Exception as exc:
        destination.unlink(missing_ok=True)
        flash(f'That file could not be read as an Excel workbook: {exc}', 'danger')
        return redirect(url_for('isms.ledgers'))

    _log_ledger_action('isms_ledger_template_upload', f'Uploaded ledger template {filename}')
    flash(f'Ledger template {filename} uploaded. It is now the active template.', 'success')
    return redirect(url_for('isms.ledgers'))


@bp.route('/isms/ledgers/export')
@login_required
@admin_required
def export_ledgers():
    try:
        output, summary = isms_ledger_service.build_ledger_workbook()
    except FileNotFoundError as exc:
        flash(str(exc), 'danger')
        return redirect(url_for('isms.ledgers'))

    filled = ', '.join(
        f"{name} {stats.get('total', 0)}" for name, stats in summary['sheets'].items()
    )
    _log_ledger_action(
        'isms_ledger_export',
        f"Generated ISMS management ledgers from template {summary['template']} ({filled})",
    )

    response = send_file(
        output,
        mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
        as_attachment=True,
        download_name=isms_ledger_service.export_filename(),
    )
    response.headers['Cache-Control'] = 'no-store, no-cache, must-revalidate, max-age=0'
    response.headers['Pragma'] = 'no-cache'
    response.headers['Expires'] = '0'
    return response
