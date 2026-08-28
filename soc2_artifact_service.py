import html
import io
import os
import re
from collections import Counter
from datetime import datetime, timedelta

from docx import Document
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer

from extensions import db
from models import Asset, Employee, SOC2PhishingCampaign, SOC2PhishingResult, SOC2ReadinessItem, SOC2SecurityTrainingRecord, SOC2Vendor, SystemDescription


SYSTEM_DESCRIPTION_SOURCE = '/var/www/tracker/content/isms/incoming/systemdescriptions.md'
QUARTERLY_TRAINING_DAYS = 90
QUALIFYING_TRAINING_STATUSES = {'Completed', 'Passed', 'Attended'}
SYSTEM_DESCRIPTION_EXCLUDED_HEADINGS = {
    'overview of operations',
    'system changes during the period',
    'optional',
}
SYSTEM_DESCRIPTION_SECTION_REPLACEMENTS = {
    'overview of the system': [
        re.compile(r'^.*?(Cirque Corporation provides capacitive touch products.*)$', re.IGNORECASE | re.DOTALL),
    ],
    'key features of the [system name] system': [
        re.compile(r'^.*?(Cirque[’\']s Touch IC platform.*)$', re.IGNORECASE | re.DOTALL),
    ],
}
SYSTEM_DESCRIPTION_LINE_PATTERNS = [
    re.compile(r'^important:?$', re.IGNORECASE),
    re.compile(r'^document revision checklist:?$', re.IGNORECASE),
    re.compile(r'^carefully revise this template', re.IGNORECASE),
    re.compile(r'^update \[company\]', re.IGNORECASE),
    re.compile(r'^update \[system name\]', re.IGNORECASE),
    re.compile(r'^replace \[vp of engineering\]', re.IGNORECASE),
    re.compile(r'^do not change the formatting', re.IGNORECASE),
    re.compile(r'^note: the template includes formatting', re.IGNORECASE),
    re.compile(r'^red text denotes', re.IGNORECASE),
    re.compile(r'^purple text denotes', re.IGNORECASE),
    re.compile(r'^black text may be edited', re.IGNORECASE),
    re.compile(r'^all content in the system description is fair game', re.IGNORECASE),
    re.compile(r'^for a type 1 report, please omit', re.IGNORECASE),
    re.compile(r'^write in the third person', re.IGNORECASE),
    re.compile(r'^remove the table of contents', re.IGNORECASE),
    re.compile(r'^attach an editable or \.docx copy', re.IGNORECASE),
    re.compile(r'^more guidance/information here\.?$', re.IGNORECASE),
    re.compile(r'^\[optional\]$', re.IGNORECASE),
    re.compile(r'^select all applicable:?$', re.IGNORECASE),
    re.compile(r'^provide a brief summary', re.IGNORECASE),
    re.compile(r'^provide a list of key features', re.IGNORECASE),
    re.compile(r'^this section generally contains', re.IGNORECASE),
    re.compile(r'^this document comprises section iii', re.IGNORECASE),
    re.compile(r'^if a customer/client wishes to review', re.IGNORECASE),
    re.compile(r'^ultimately, it is your call', re.IGNORECASE),
    re.compile(r'^select a descriptive name for your product', re.IGNORECASE),
    re.compile(r'^provide a high-level description of the services', re.IGNORECASE),
    re.compile(r'^do not include one-off solutions', re.IGNORECASE),
    re.compile(r'^keep the list concise and focus', re.IGNORECASE),
    re.compile(r'^it is important to give your system', re.IGNORECASE),
    re.compile(r'^for example,', re.IGNORECASE),
    re.compile(r'^update these paragraphs:?$', re.IGNORECASE),
    re.compile(r'^remove rows for the tscs', re.IGNORECASE),
    re.compile(r'^update this paragraph:?$', re.IGNORECASE),
    re.compile(r'^for the company background include:?$', re.IGNORECASE),
    re.compile(r'^1\) data flow:', re.IGNORECASE),
    re.compile(r'^third party access \(to data\):', re.IGNORECASE),
    re.compile(r'^\(optional\) if you are familiar', re.IGNORECASE),
]
SYSTEM_DESCRIPTION_BULLET_PATTERNS = [
    re.compile(r'^the company mission statement\.?$', re.IGNORECASE),
    re.compile(r'^the location/s of the company\.?$', re.IGNORECASE),
    re.compile(r'^industries that the company serves\.?$', re.IGNORECASE),
    re.compile(r'^a brief history of the company\.?$', re.IGNORECASE),
    re.compile(r'^customers interested in using our products$', re.IGNORECASE),
    re.compile(r'^existing users of our products for all or a portion of the reporting period$', re.IGNORECASE),
    re.compile(r'^internal personnel$', re.IGNORECASE),
    re.compile(r'^our service auditor$', re.IGNORECASE),
    re.compile(r'^the service auditors of the entities that interact with our products$', re.IGNORECASE),
    re.compile(r'^other business partners that interact with our products$', re.IGNORECASE),
    re.compile(r'^regulatory authorities, if needed\.?$', re.IGNORECASE),
]


def _normalize_heading(value):
    return re.sub(r'[^a-z0-9]+', ' ', (value or '').lower()).strip()


def _markdown_to_docx_paragraphs(document, markdown_content):
    for raw_line in markdown_content.splitlines():
        line = raw_line.rstrip()
        if not line:
            document.add_paragraph('')
            continue
        if line.startswith('### '):
            document.add_heading(line[4:], level=3)
        elif line.startswith('## '):
            document.add_heading(line[3:], level=2)
        elif line.startswith('# '):
            document.add_heading(line[2:], level=1)
        elif line.startswith('- '):
            document.add_paragraph(line[2:], style='List Bullet')
        else:
            document.add_paragraph(line)


def get_training_completion_snapshot(cycle_days=QUARTERLY_TRAINING_DAYS):
    # Only our own people are in scope for security-awareness training.
    # Partner staff (badge access only) have no AD account, so an ad_enabled
    # test alone would have counted them and reported them overdue.
    employees = Employee.reportable().order_by(Employee.name.asc()).all()
    active_employees = [employee for employee in employees if employee.ad_enabled is not False]
    cutoff_date = datetime.utcnow().date() - timedelta(days=cycle_days)
    records = SOC2SecurityTrainingRecord.query.order_by(
        SOC2SecurityTrainingRecord.training_date.desc(),
        SOC2SecurityTrainingRecord.id.desc(),
    ).all()

    latest_completed_by_employee = {}
    for record in records:
        if record.employee_id is None:
            continue
        if record.employee_id in latest_completed_by_employee:
            continue
        if record.completion_status not in QUALIFYING_TRAINING_STATUSES:
            continue
        latest_completed_by_employee[record.employee_id] = record

    completed_current_cycle = 0
    overdue_employees = []
    for employee in active_employees:
        latest_record = latest_completed_by_employee.get(employee.id)
        if latest_record and latest_record.training_date and latest_record.training_date >= cutoff_date:
            completed_current_cycle += 1
        else:
            overdue_employees.append(employee)

    return {
        'employees': employees,
        'active_employees': active_employees,
        'records': records,
        'completed_current_cycle': completed_current_cycle,
        'overdue_employees': overdue_employees,
        'cutoff_date': cutoff_date,
    }


def import_system_description_from_markdown(updated_by='system_import'):
    if not os.path.exists(SYSTEM_DESCRIPTION_SOURCE):
        raise FileNotFoundError(SYSTEM_DESCRIPTION_SOURCE)

    with open(SYSTEM_DESCRIPTION_SOURCE, 'r', encoding='utf-8') as handle:
        raw_content = handle.read().replace('\r\n', '\n')

    sections = SystemDescription.query.order_by(SystemDescription.section_order.asc(), SystemDescription.id.asc()).all()
    if not sections:
        section = SystemDescription(
            section_title='Imported System Description',
            section_level=1,
            section_order=0,
            category='general',
            content=raw_content.strip(),
            template_content=None,
            auto_populated=False,
            updated_by=updated_by,
        )
        db.session.add(section)
        db.session.commit()
        return {'sections': 1, 'updated': 1, 'matched': 1}

    section_by_heading = {_normalize_heading(section.section_title): section for section in sections}
    collected = {section.id: [] for section in sections}
    matched_sections = set()
    current_section = None
    intro_lines = []

    for raw_line in raw_content.splitlines():
        stripped = raw_line.strip().lstrip('#').strip()
        normalized = _normalize_heading(stripped)
        if normalized and normalized in section_by_heading:
            current_section = section_by_heading[normalized]
            matched_sections.add(current_section.id)
            continue
        if current_section is None:
            intro_lines.append(raw_line.rstrip())
        else:
            collected[current_section.id].append(raw_line.rstrip())

    updated_count = 0
    for section in sections:
        if _normalize_heading(section.section_title) in SYSTEM_DESCRIPTION_EXCLUDED_HEADINGS:
            content = ''
        else:
            content = _clean_system_description_content(section.section_title, collected[section.id])
        if section.content != content:
            section.content = content
            section.updated_by = updated_by
            updated_count += 1

    db.session.commit()
    return {'sections': len(sections), 'updated': updated_count, 'matched': len(matched_sections)}


def build_system_description_markdown():
    sections = SystemDescription.query.order_by(SystemDescription.section_order.asc(), SystemDescription.id.asc()).all()
    assets = Asset.query.all()
    vendors = SOC2Vendor.query.filter_by(is_active=True).all()
    readiness_items = SOC2ReadinessItem.query.filter_by(is_active=True).all()
    training_snapshot = get_training_completion_snapshot()

    lines = [
        '# Tracker System Description',
        '',
        f'Generated: {datetime.utcnow().strftime("%Y-%m-%d %H:%M:%S UTC")}',
        '',
    ]

    for section in sections:
        level = min(max((section.section_level or 1) + 1, 2), 6)
        content = (section.content or section.template_content or '').strip()
        if not content:
            continue
        lines.extend([f'{'#' * level} {section.section_title}', '', content, ''])

    asset_categories = Counter(asset.category or 'Uncategorized' for asset in assets)
    departments = Counter(employee.department or 'Unassigned' for employee in training_snapshot['active_employees'])
    open_readiness = sum(1 for item in readiness_items if item.status in {'Not In Place', 'Partially In Place', 'Open', 'Blocked'})

    lines.extend([
        '## Tracker Data Appendix',
        '',
        f'- Assets tracked: {len(assets)}',
        f'- Active personnel tracked: {len(training_snapshot["active_employees"])}',
        f'- Active vendors tracked: {len(vendors)}',
        f'- Open readiness items: {open_readiness}',
        f'- Security training records: {len(training_snapshot["records"])}',
        f'- Personnel current on quarterly security training: {training_snapshot["completed_current_cycle"]}',
        '',
        '### Asset Categories',
        '',
    ])

    for category, count in sorted(asset_categories.items()):
        lines.append(f'- {category}: {count}')

    lines.extend(['', '### Departments', ''])
    for department, count in sorted(departments.items()):
        lines.append(f'- {department}: {count}')

    return '\n'.join(lines).strip() + '\n'


def get_phishing_campaign_snapshot():
    campaigns = SOC2PhishingCampaign.query.order_by(SOC2PhishingCampaign.campaign_date.desc(), SOC2PhishingCampaign.id.desc()).all()
    results = SOC2PhishingResult.query.order_by(SOC2PhishingResult.training_completed_on.desc().nullslast(), SOC2PhishingResult.id.desc()).all()
    return {
        'campaigns': campaigns,
        'results': results,
        'employees_trained': sum(1 for result in results if result.training_completed),
        'clicks': sum(1 for result in results if result.clicked),
        'reports': sum(1 for result in results if result.reported),
    }


def build_system_description_docx():
    document = Document()
    _markdown_to_docx_paragraphs(document, build_system_description_markdown())
    buffer = io.BytesIO()
    document.save(buffer)
    buffer.seek(0)
    return buffer


def build_system_description_pdf():
    markdown_content = build_system_description_markdown()
    buffer = io.BytesIO()
    document = SimpleDocTemplate(
        buffer,
        pagesize=letter,
        leftMargin=0.75 * inch,
        rightMargin=0.75 * inch,
        topMargin=0.75 * inch,
        bottomMargin=0.75 * inch,
    )
    styles = getSampleStyleSheet()
    heading_style = ParagraphStyle('SystemHeading', parent=styles['Heading1'], spaceAfter=8)
    subheading_style = ParagraphStyle('SystemSubHeading', parent=styles['Heading2'], spaceAfter=6)
    body_style = ParagraphStyle('SystemBody', parent=styles['BodyText'], leading=14, spaceAfter=6)

    story = []
    for raw_line in markdown_content.splitlines():
        line = raw_line.strip()
        if not line:
            story.append(Spacer(1, 0.1 * inch))
            continue
        if line.startswith('# '):
            story.append(Paragraph(html.escape(line[2:]), heading_style))
        elif line.startswith('## '):
            story.append(Paragraph(html.escape(line[3:]), subheading_style))
        elif line.startswith('### '):
            story.append(Paragraph(html.escape(line[4:]), subheading_style))
        elif line.startswith('- '):
            story.append(Paragraph(f'• {html.escape(line[2:])}', body_style))
        else:
            story.append(Paragraph(html.escape(line), body_style))

    document.build(story)
    buffer.seek(0)
    return buffer


def _should_drop_system_description_line(line):
    normalized = _normalize_heading(line)
    if not normalized:
        return False
    if '[' in line and ']' in line:
        return True
    if any(pattern.match(line.strip()) for pattern in SYSTEM_DESCRIPTION_LINE_PATTERNS):
        return True
    bullet_candidate = line.strip().lstrip('-').strip()
    if any(pattern.match(bullet_candidate) for pattern in SYSTEM_DESCRIPTION_BULLET_PATTERNS):
        return True
    return False


def _clean_system_description_content(section_title, lines):
    cleaned = []
    previous_blank = True
    normalized_title = _normalize_heading(section_title)
    for raw_line in lines:
        line = raw_line.rstrip()
        stripped = line.strip()
        if not stripped:
            if not previous_blank:
                cleaned.append('')
            previous_blank = True
            continue
        if _normalize_heading(stripped) == normalized_title:
            continue
        for pattern in SYSTEM_DESCRIPTION_SECTION_REPLACEMENTS.get(normalized_title, []):
            match = pattern.match(stripped)
            if match:
                stripped = match.group(1).strip()
                line = stripped
                break
        if _should_drop_system_description_line(stripped):
            continue
        cleaned.append(line)
        previous_blank = False

    while cleaned and not cleaned[-1].strip():
        cleaned.pop()
    content = '\n'.join(cleaned).strip()
    for pattern in SYSTEM_DESCRIPTION_SECTION_REPLACEMENTS.get(normalized_title, []):
        match = pattern.match(content)
        if match:
            content = match.group(1).strip()
    return content


def enrich_system_description_sections(updated_by='system_enrichment', commit=True):
    sections = SystemDescription.query.order_by(SystemDescription.section_order.asc(), SystemDescription.id.asc()).all()
    assets = Asset.query.all()
    vendors = SOC2Vendor.query.filter_by(is_active=True).all()
    employees = [employee for employee in Employee.query.order_by(Employee.name.asc()).all() if employee.ad_enabled is not False]
    readiness_items = SOC2ReadinessItem.query.filter_by(is_active=True).all()

    asset_categories = Counter(asset.category or 'Uncategorized' for asset in assets)
    software_components = Counter((asset.os_version or asset.model or asset.category or 'Unspecified') for asset in assets)
    departments = Counter(employee.department or 'Unassigned' for employee in employees)

    section_content = {
        'overview of operations': (
            'Cirque Corporation operates a centralized information security program supporting product development, manufacturing support, and corporate operations across its in-scope environment. '
            'Tracker acts as the operational system of record for assets, personnel alignment, vendors, governance actions, training, phishing outcomes, and audit evidence used during SOC 2 preparation.'
        ),
        'key features of the [system name] system': '\n'.join([
            '- Centralized asset inventory with ownership, location, and lifecycle tracking.',
            '- Directory-aligned personnel records used for audit scoping and evidence correlation.',
            '- Native readiness, internal audit, vendor, management review, acknowledgement, training, and phishing workflows.',
            '- Evidence-pack generation for point-in-time SOC 2 Type 1 reporting.',
        ]),
        'software': '\n'.join([
            'Tracker relies on the following managed software or operating-system components identified from the current inventory:',
            '',
            *[f'- {component}: {count} tracked assets' for component, count in software_components.most_common(12)],
        ]),
        'relevant aspects of the control environment risk assessment information and communications and monitoring': (
            'The system is governed through formal management oversight, documented policies, recurring training, monitoring workflows, readiness tracking, and evidence-backed operational reviews. '
            'Tracker consolidates information and communications needed to support control monitoring, issue remediation, and audit preparation.'
        ),
        'controls': (
            f'Tracker currently maintains {len(readiness_items)} active readiness items and supporting evidence from internal audits, management reviews, vendor reviews, training records, policy acknowledgements, and phishing outcomes. '
            'These control activities support the operating effectiveness narrative described throughout the system description.'
        ),
        'boundaries of the system': (
            'The in-scope system includes managed corporate assets, identity-linked user records, governance workflows, and supporting evidence maintained in Tracker. '
            'External customer environments and third-party service-provider internal controls remain outside direct operational control except where Cirque manages risk through vendor oversight and contractual commitments.'
        ),
    }

    updated = 0
    for section in sections:
        normalized_title = _normalize_heading(section.section_title)
        if normalized_title == 'people' and not (section.content or '').strip():
            section.content = '\n'.join([
                f'Cirque currently tracks {len(employees)} active personnel within the in-scope environment.',
                '',
                *[f'- {department}: {count} active personnel' for department, count in sorted(departments.items())],
            ])
            section.updated_by = updated_by
            updated += 1
            continue
        if normalized_title == 'third party access' and not (section.content or '').strip():
            section.content = '\n'.join([
                'Third-party access is managed through vendor oversight, defined service scope, and evidence-backed periodic reviews.',
                '',
                *[f'- {vendor.vendor_name}: {vendor.service_description or vendor.vendor_type or "Managed service"}' for vendor in vendors[:12]],
            ])
            section.updated_by = updated_by
            updated += 1
            continue
        if normalized_title == 'infrastructure' and not (section.content or '').strip():
            section.content = '\n'.join([
                'Infrastructure supporting the in-scope system is represented in the managed asset inventory and includes endpoint, server, and supporting platform components.',
                '',
                *[f'- {category}: {count} tracked assets' for category, count in sorted(asset_categories.items())],
            ])
            section.updated_by = updated_by
            updated += 1
            continue
        if normalized_title in section_content and not (section.content or '').strip():
            section.content = section_content[normalized_title]
            section.updated_by = updated_by
            updated += 1

    if commit:
        db.session.commit()
    return {'updated': updated}